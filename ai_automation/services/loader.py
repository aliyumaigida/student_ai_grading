# grading/services/loader.py
import os
import re
import json
import pandas as pd
import csv
from docx import Document
import pdfplumber
import tempfile
from PIL import Image
import pytesseract

# Set Tesseract path
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


def load_marking_scheme(file):
    """
    Load marking scheme from csv, json, excel, docx, txt, or pdf.
    Returns a list of question dicts: {question, marking_scheme, max_marks}
    """
    ext = get_extension(file)

    if ext == ".csv":
        return load_csv(file)
    elif ext == ".json":
        return load_json(file)
    elif ext in [".xls", ".xlsx"]:
        return load_excel(file)
    elif ext == ".docx":
        return load_docx(file)
    elif ext in [".txt", ".pdf"]:
        text = extract_text(file)
        if not text:
            raise Exception("Failed to extract text from file")
        return parse_questions_from_text(text)
    else:
        raise Exception("Unsupported format: " + ext)


# ---------------- Helper functions ----------------

def get_extension(file):
    """Return lowercase extension including dot, e.g., '.csv'"""
    if hasattr(file, "name"):
        return os.path.splitext(file.name)[1].lower()
    elif isinstance(file, str) and os.path.exists(file):
        return os.path.splitext(file)[1].lower()
    else:
        raise ValueError("Invalid file: must be path or file-like object")


def load_csv(file):
    if hasattr(file, "read"):
        file.seek(0)
        return list(csv.DictReader(file.read().decode("utf-8").splitlines()))
    else:
        with open(file, encoding="utf-8") as f:
            return list(csv.DictReader(f))


def load_json(file):
    if hasattr(file, "read"):
        file.seek(0)
        return json.load(file)
    else:
        with open(file, encoding="utf-8") as f:
            return json.load(f)


def load_excel(file):
    if hasattr(file, "read"):
        return pd.read_excel(file).to_dict("records")
    else:
        return pd.read_excel(file).to_dict("records")


def load_docx(file):
    if hasattr(file, "read") or (isinstance(file, str) and os.path.exists(file)):
        doc = Document(file)
    else:
        raise ValueError("Invalid file for DOCX")

    return parse_questions_from_text("\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()]))


def extract_text(file):
    """
    Extract text from txt, docx, or pdf (with OCR fallback).
    Returns plain text string.
    """
    ext = get_extension(file)

    if ext == ".txt":
        if hasattr(file, "read"):
            return file.read().decode("utf-8", errors="ignore")
        else:
            with open(file, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

    elif ext == ".docx":
        doc = Document(file)
        return "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

    elif ext == ".pdf":
        text = ""
        try:
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    # ✅ Use OCR only if text is empty or very short
                    if page_text and len(page_text.strip()) > 10:
                        text += page_text + "\n"
                    else:
                        # OCR fallback
                        image = page.to_image(resolution=300).original
                        with tempfile.NamedTemporaryFile(suffix=".png") as tmp:
                            image.save(tmp.name, format="PNG")
                            ocr_text = pytesseract.image_to_string(Image.open(tmp.name))
                        text += ocr_text + "\n"
        except Exception as e:
            print("❌ PDF extraction error:", e)
            return ""

        # ✅ Remove duplicate lines while preserving order
        lines = text.splitlines()
        seen = set()
        unique_lines = []
        for line in lines:
            clean_line = line.strip()
            if clean_line and clean_line not in seen:
                unique_lines.append(clean_line)
                seen.add(clean_line)
        return "\n".join(unique_lines)

    else:
        raise Exception("Unsupported file format: " + ext)


def parse_questions_from_text(text):
    """
    Parses text for questions using 'Question <number>' pattern.
    Returns a list of dicts: {question, marking_scheme, max_marks}
    """
    q_pattern = re.compile(r"^Question\s+\d+", re.I)
    marks_pattern = re.compile(r"\((\d+)")

    exam = []
    current = None
    lines = []

    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue

        if q_pattern.match(line):
            if current:
                marks = sum(map(int, marks_pattern.findall("\n".join(lines))))
                exam.append({
                    "question": current,
                    "max_marks": marks or 100,
                    "marking_scheme": "\n".join(lines)
                })
            current = line
            lines = []
        else:
            lines.append(line)

    # Add the last question
    if current:
        marks = sum(map(int, marks_pattern.findall("\n".join(lines))))
        exam.append({
            "question": current,
            "max_marks": marks or 100,
            "marking_scheme": "\n".join(lines)
        })

    # Fallback: if no question headers detected, treat each line as a question
    if not exam:
        exam = [{"question": line, "marking_scheme": line, "max_marks": 100} for line in text.splitlines() if line.strip()]

    return exam
