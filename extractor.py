import pdfplumber
from docx import Document
import os
import re

# -------------------------
# Text Extraction Helpers
# -------------------------
def clean_text(text: str) -> str:
    """Clean extracted text"""
    if not text:
        return ""
    # Remove extra whitespace and repeated newlines
    text = re.sub(r"\n\s*\n", "\n", text)
    return text.strip()


def extract_pdf(file_path: str) -> str:
    text = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            content = page.extract_text()
            if content:
                text.append(content)
    return "\n".join(text)


def extract_docx(file_path: str) -> str:
    """Extract text from DOCX (student submissions)"""
    doc = Document(file_path)
    return "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])


def extract_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text(file_path: str) -> str:
    """Main extractor function for student files"""
    if not os.path.exists(file_path):
        print("❌ File does not exist")
        return None

    try:
        file_path = file_path.strip('"')  # remove quotes if pasted path
        ext = file_path.lower().split(".")[-1]

        print(f"📄 Extracting from: {file_path}")

        if ext == "pdf":
            text = extract_pdf(file_path)

        elif ext == "docx":
            text = extract_docx(file_path)

        elif ext == "txt":
            text = extract_txt(file_path)

        else:
            print("❌ Unsupported file format")
            return None

        text = clean_text(text)
        if not text:
            print("⚠️ No text found in file")
            return None

        return text

    except Exception as e:
        print(f"❌ Extraction error: {e}")
        return None


# -------------------------
# DOCX Marking Scheme Loader
# -------------------------
def load_marking_scheme_docx(file_path: str):
    """
    Parses DOCX marking scheme with sub-questions like 1a, 1b, 2a, 2b.
    Returns list of dicts:
    [{question: str, max_marks: int, marking_scheme: str}, ...]
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"❌ Marking scheme file not found: {file_path}")

    doc = Document(file_path)
    exam = []

    current_question = ""
    current_answer = ""
    current_marks = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Match sub-question like '1a.' or '2b.'
        q_match = re.match(r"^(\d+[a-zA-Z]?)\.\s*(.*)", text)
        if q_match:
            # Save previous question
            if current_question:
                exam.append({
                    "question": current_question,
                    "max_marks": current_marks if current_marks > 0 else 1,
                    "marking_scheme": current_answer.strip()
                })

            # Start new question
            current_question = q_match.group(0)  # '1a. Define Database'
            current_answer = q_match.group(2) or ""

            # Extract marks in parentheses (e.g., (6 marks))
            marks_match = re.findall(r"\((\d+)\s*marks?\)", text, re.IGNORECASE)
            current_marks = sum(int(m) for m in marks_match) if marks_match else 0

        else:
            current_answer += "\n" + text

    # Append last question
    if current_question:
        exam.append({
            "question": current_question,
            "max_marks": current_marks if current_marks > 0 else 1,
            "marking_scheme": current_answer.strip()
        })

    return exam
